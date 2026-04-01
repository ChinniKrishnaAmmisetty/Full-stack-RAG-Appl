"""Document text extraction, validation, and chunking service."""

import os
import io
import csv
import logging
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Magic bytes for file type validation
FILE_SIGNATURES = {
    "pdf": [b"%PDF"],
    "docx": [b"PK\x03\x04"],  # DOCX is a ZIP archive
    "xlsx": [b"PK\x03\x04"],  # XLSX is also a ZIP archive
}


def validate_file_content(file_path: str, file_type: str) -> bool:
    """Validate that the file content matches its claimed type using magic bytes.

    Returns True if the file is valid, False if suspicious.
    Text-based formats (txt, csv, md) are always considered valid.
    """
    # Text-based formats can't be validated by magic bytes
    if file_type in ("txt", "csv", "md"):
        return True

    signatures = FILE_SIGNATURES.get(file_type)
    if not signatures:
        return False

    try:
        with open(file_path, "rb") as f:
            header = f.read(16)
        return any(header.startswith(sig) for sig in signatures)
    except Exception:
        return False


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text content from a file based on its type."""
    file_type = file_type.lower()

    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "txt": _extract_txt,
        "csv": _extract_csv,
        "xlsx": _extract_xlsx,
        "md": _extract_txt,  # Markdown is plain text
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(file_path)


def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    """Extract text from a DOCX file, including tables."""
    doc = DocxDocument(file_path)
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_txt(file_path: str) -> str:
    """Extract text from a TXT or Markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_csv(file_path: str) -> str:
    """Extract text from a CSV file, converting each row to a readable line."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            # Sniff the dialect to handle different CSV formats
            sample = f.read(8192)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel

            reader = csv.reader(f, dialect)
            rows = list(reader)

        if not rows:
            return ""

        # Use first row as headers if it looks like a header row
        headers = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else rows

        lines = []
        for row in data_rows:
            # Format each row as "header1: value1, header2: value2, ..."
            if len(row) == len(headers) and data_rows != rows:
                line = ", ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
            else:
                line = ", ".join(v for v in row if v.strip())
            if line.strip():
                lines.append(line)

        return "\n".join(lines)
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        raise ValueError(f"Failed to extract text from CSV: {e}")


def _extract_xlsx(file_path: str) -> str:
    """Extract text from an Excel file, processing all sheets."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        all_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = list(sheet.iter_rows(values_only=True))

            if not rows:
                continue

            all_text.append(f"\n## Sheet: {sheet_name}\n")

            # Use first row as headers
            headers = [str(cell) if cell is not None else "" for cell in rows[0]]
            data_rows = rows[1:] if len(rows) > 1 else rows

            for row in data_rows:
                values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in values):
                    if len(values) == len(headers) and data_rows != rows:
                        line = ", ".join(f"{h}: {v}" for h, v in zip(headers, values) if v.strip())
                    else:
                        line = ", ".join(v for v in values if v.strip())
                    if line.strip():
                        all_text.append(line)

        wb.close()
        return "\n".join(all_text)
    except ImportError:
        raise ValueError("openpyxl is required for Excel file support. Install it with: pip install openpyxl")
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        raise ValueError(f"Failed to extract text from Excel file: {e}")


# ─── Chunking ─────────────────────────────────────────────────

# Separators ordered from most to least significant for recursive splitting
_SEPARATORS = ["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " "]


def split_text_into_chunks(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks using recursive character-based splitting.

    This is much better than naive word splitting because it:
    - Preserves paragraph and sentence boundaries where possible
    - Falls back to smaller separators only when needed
    - Keeps semantically coherent units together

    Args:
        text: The full text to split.
        chunk_size: Target number of words per chunk.
        chunk_overlap: Number of overlapping words between consecutive chunks.

    Returns:
        A list of text chunks.
    """
    if not text or not text.strip():
        return []

    # First, split recursively by structural separators
    raw_segments = _recursive_split(text.strip(), _SEPARATORS, chunk_size)

    # Then merge small segments and create overlapping windows
    chunks = _merge_with_overlap(raw_segments, chunk_size, chunk_overlap)

    return [c for c in chunks if c.strip()]


def _recursive_split(text: str, separators: list[str], chunk_size: int) -> list[str]:
    """Recursively split text using a hierarchy of separators."""
    # If text is already small enough, return it as-is
    word_count = len(text.split())
    if word_count <= chunk_size:
        return [text]

    # Try each separator in order (most significant first)
    for sep in separators:
        if sep in text:
            parts = text.split(sep)
            # Only use this separator if it actually creates multiple parts
            if len(parts) > 1:
                result = []
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    # Add the separator back (except for space) for readability
                    if sep not in (" ", ", "):
                        part_with_sep = part + (sep.strip() if sep.strip() else "")
                    else:
                        part_with_sep = part

                    # If this part is still too large, recurse with remaining separators
                    if len(part_with_sep.split()) > chunk_size:
                        remaining_seps = separators[separators.index(sep) + 1:]
                        if remaining_seps:
                            result.extend(_recursive_split(part_with_sep, remaining_seps, chunk_size))
                        else:
                            # Last resort: hard word split
                            result.extend(_hard_word_split(part_with_sep, chunk_size))
                    else:
                        result.append(part_with_sep)
                return result

    # No separator worked — hard split by word count
    return _hard_word_split(text, chunk_size)


def _hard_word_split(text: str, chunk_size: int) -> list[str]:
    """Last-resort splitting by word count when no separators work."""
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
    return chunks


def _merge_with_overlap(segments: list[str], chunk_size: int, chunk_overlap: int) -> list[str]:
    """Merge small segments together and create overlapping chunks."""
    if not segments:
        return []

    chunks = []
    current_words = []

    for segment in segments:
        seg_words = segment.split()

        # If adding this segment would exceed chunk size, finalize current chunk
        if current_words and len(current_words) + len(seg_words) > chunk_size:
            chunks.append(" ".join(current_words))
            # Keep overlap words from the end of current chunk
            if chunk_overlap > 0 and len(current_words) > chunk_overlap:
                current_words = current_words[-chunk_overlap:]
            else:
                current_words = []

        current_words.extend(seg_words)

    # Don't forget the last chunk
    if current_words:
        chunks.append(" ".join(current_words))

    return chunks

